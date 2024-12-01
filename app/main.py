from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, status, Body
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy.orm import Session
from pydantic import BaseModel, Field
from docx import Document
from models import SessionLocal, Template, GeneratedDocument
import os
import shutil
from datetime import datetime
from typing import Optional, Union

# Путь к папке app
APP_DIR = os.path.dirname(os.path.abspath(__file__))

# Путь к папке templates
TEMPLATES_DIR = os.path.join(APP_DIR, "templates")

# Путь к папке documents
DOCUMENTS_DIR = os.path.join(APP_DIR, "documents")

# Создаём папки, если их нет
os.makedirs(TEMPLATES_DIR, exist_ok=True)
os.makedirs(DOCUMENTS_DIR, exist_ok=True)

app = FastAPI()

# Подключение CORS для фронтенда
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Замените на нужный URL фронтенда
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Зависимость для подключения к базе данных
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()

# --- Модели запросов ---
class DocumentData(BaseModel):
    seller_name: str
    buyer_name: str
    item: str
    price: float = Field(gt=0, description="Цена должна быть больше 0")

class LegalServicesData(BaseModel):
    contract_date: str
    lawyer_name: str
    client_name: str
    client_passport_series: str
    client_passport_number: str
    client_passport_issued_by: str
    client_passport_issued_date: str
    client_address: str

# --- Маршрут для загрузки шаблонов ---
@app.post("/upload-template/")
async def upload_template(file: UploadFile = File(...), db: Session = Depends(get_db)):
    # Генерируем путь для сохранения шаблона в папке templates
    template_path = os.path.join(TEMPLATES_DIR, file.filename)

    # Проверяем, существует ли файл
    if os.path.exists(template_path):
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Шаблон с таким именем уже существует")

    # Сохраняем файл в папку
    with open(template_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    # Создаем запись в базе данных для нового шаблона
    new_template = Template(
        name=file.filename.split(".")[0],  # Убираем расширение для названия
        description=f"Шаблон {file.filename.split('.')[0]}",  # Можно добавить описание
        file_path=template_path
    )
    
    db.add(new_template)
    db.commit()
    db.refresh(new_template)

    return {"message": "Шаблон успешно загружен", "template_id": new_template.id}

# Получение списка шаблонов
@app.get("/templates/")
async def list_templates(db: Session = Depends(get_db)):
    templates = db.query(Template).all()
    return templates

# --- Маршрут для генерации документа ---
@app.post("/generate/{template_id}/")
async def generate_document(
    template_id: int,
    data: Union[DocumentData, LegalServicesData] = Body(...),  # Используем Body для данных из тела запроса
    db: Session = Depends(get_db)
):
    template = db.query(Template).filter(Template.id == template_id).first()
    if not template:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Шаблон не найден")
    
    # Открываем шаблон
    template_path = template.file_path
    doc = Document(template_path)
    
    # Заполняем плейсхолдеры в зависимости от типа шаблона
    if template.name == "BuySellContract" and isinstance(data, DocumentData):
        for paragraph in doc.paragraphs:
            if "{{seller_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{seller_name}}", data.seller_name)
            if "{{buyer_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{buyer_name}}", data.buyer_name)
            if "{{item}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{item}}", data.item)
            if "{{price}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{price}}", str(data.price))

    elif template.name == "LegalServicesContract" and isinstance(data, LegalServicesData):
        for paragraph in doc.paragraphs:
            if "{{contract_date}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{contract_date}}", data.contract_date)
            if "{{lawyer_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{lawyer_name}}", data.lawyer_name)
            if "{{client_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_name}}", data.client_name)
            if "{{client_passport_series}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_series}}", data.client_passport_series)
            if "{{client_passport_number}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_number}}", data.client_passport_number)
            if "{{client_passport_issued_by}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_issued_by}}", data.client_passport_issued_by)
            if "{{client_passport_issued_date}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_issued_date}}", data.client_passport_issued_date)
            if "{{client_address}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_address}}", data.client_address)

    # Генерация уникального имени файла
    today = datetime.today()
    date_str = today.strftime("%y%m%d")
    existing_documents_today = db.query(GeneratedDocument).filter(
        GeneratedDocument.template_id == template.id,
        GeneratedDocument.created_at.startswith(today.strftime("%Y-%m-%d"))
    ).count()
    
    document_number = existing_documents_today + 1
    file_name = f"{date_str}-{document_number}_{template.name}.docx"
    file_path = os.path.join(DOCUMENTS_DIR, file_name)

    # Сохраняем новый документ
    doc.save(file_path)

    # Создаём запись в базе данных
    new_document = GeneratedDocument(
        template_id=template.id,
        file_path=file_path,
        created_at=datetime.now()
    )
    db.add(new_document)
    db.commit()
    db.refresh(new_document)

    # Возвращаем путь к документу для скачивания
    return {"message": "Документ создан", "document_id": new_document.id, "file_path": f"/documents/{new_document.id}"}

# Маршрут для скачивания документа
@app.get("/documents/{document_id}")
async def download_document(document_id: int, db: Session = Depends(get_db)):
    document = db.query(GeneratedDocument).filter(GeneratedDocument.id == document_id).first()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Документ не найден")
    
    # Проверяем, существует ли файл по указанному пути
    if not os.path.exists(document.file_path):
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Файл не найден на сервере")

    # Возвращаем файл пользователю
    return FileResponse(document.file_path, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", filename=document.file_path.split("/")[-1])

# Маршрут для редактирования документа
@app.post("/edit-document/{document_id}/")
async def edit_document(
    document_id: int,
    data: Union[DocumentData, LegalServicesData] = Body(...),
    db: Session = Depends(get_db)
):
    document = db.query(GeneratedDocument).filter(GeneratedDocument.id == document_id).first()
    if not document:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Документ не найден")
    
    # Открываем документ для редактирования
    doc = Document(document.file_path)
    
    # Вносим изменения, аналогично генерации документа
    if document.template.name == "BuySellContract" and isinstance(data, DocumentData):
        for paragraph in doc.paragraphs:
            if "{{seller_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{seller_name}}", data.seller_name)
            if "{{buyer_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{buyer_name}}", data.buyer_name)
            if "{{item}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{item}}", data.item)
            if "{{price}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{price}}", str(data.price))
    
    elif document.template.name == "LegalServicesContract" and isinstance(data, LegalServicesData):
        for paragraph in doc.paragraphs:
            if "{{contract_date}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{contract_date}}", data.contract_date)
            if "{{lawyer_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{lawyer_name}}", data.lawyer_name)
            if "{{client_name}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_name}}", data.client_name)
            if "{{client_passport_series}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_series}}", data.client_passport_series)
            if "{{client_passport_number}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_number}}", data.client_passport_number)
            if "{{client_passport_issued_by}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_issued_by}}", data.client_passport_issued_by)
            if "{{client_passport_issued_date}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_passport_issued_date}}", data.client_passport_issued_date)
            if "{{client_address}}" in paragraph.text:
                paragraph.text = paragraph.text.replace("{{client_address}}", data.client_address)

    # Сохраняем изменения
    doc.save(document.file_path)

    return {"message": "Документ успешно отредактирован", "document_id": document.id}
